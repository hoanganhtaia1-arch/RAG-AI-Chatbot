import os
import multiprocessing
import re
from pathlib import Path
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core import Document
from llama_index.core.schema import MediaResource

# Workaround for environments without /dev/shm (common in sandboxes/containers)
# where Intel OpenMP can crash with "Can't open SHM2".
os.environ.setdefault("KMP_USE_SHM", "0")
os.environ.setdefault("TOKENIZERS_PARALLELISM", "false")

# ── Embedding model (local, không gửi dữ liệu ra ngoài) ───────────────────────
# NOTE: Must match Qdrant vector dim (default 1024 in `vector_db.py`).
_MODEL_NAME = os.getenv("EMBED_MODEL", "intfloat/multilingual-e5-large")
_model = None

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200


# ── Constants ─────────────────────────────────────────────────────────────────
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".tif", ".webp"}

# ── Public API ─────────────────────────────────────────────────────────────────

def _ocr_extract_from_image(image_path_or_pil) -> str:
    """Helper để trích xuất text từ 1 ảnh hoặc đối tượng PIL bằng Tesseract."""
    import pytesseract
    try:
        text = pytesseract.image_to_string(image_path_or_pil, lang="vie+eng")
        return text.strip()
    except Exception as e:
        print(f"[data_loader] OCR Error: {e}")
        return ""

def load_and_chunk_pdf(path: str) -> list[dict]:
    """
    Load file PDF / DOCX / HTML / Ảnh bằng PyMuPDF4LLM kết hợp Tesseract OCR.
    Hỗ trợ bóc tách Bảng Biểu và Bố Cục Markdown, tự động fallback sang OCR nếu là tài liệu scan.
    """
    path_obj = Path(path)

    if not path_obj.exists():
        raise FileNotFoundError(f"File không tồn tại: {path}")
    if path_obj.stat().st_size == 0:
        raise ValueError(f"File rỗng (0 byte): {path}")

    print(f"[data_loader] Đang phân tích cấu trúc file: {path_obj.name}...")
    file_type = path_obj.suffix.lower().lstrip(".") or "unknown"
    markdown_text = ""

    try:
        if file_type == "docx":
            import docx
            doc = docx.Document(path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))
            markdown_text = "\n\n".join(full_text)
        elif path_obj.suffix.lower() in IMAGE_EXTENSIONS:
            print(f"[data_loader] Nhận diện định dạng ảnh -> Sử dụng Tesseract OCR...")
            markdown_text = _ocr_extract_from_image(path)
        else:
            import pymupdf4llm
            # Tự động parse layout & table thành chuỗi Markdown hoàn chỉnh
            markdown_text = pymupdf4llm.to_markdown(path)
            
            # Kiểm tra xem có phải PDF scan không (ít text)
            if len(markdown_text.strip()) < 50:
                print(f"[data_loader] Nội dung quá ngắn ({len(markdown_text)} ký tự), khả năng là PDF scan -> Thử OCR...")
                from pdf2image import convert_from_path
                try:
                    images = convert_from_path(path)
                    ocr_texts = []
                    for i, img in enumerate(images):
                        page_text = _ocr_extract_from_image(img)
                        if page_text:
                            ocr_texts.append(f"--- Trang {i+1} ---\n{page_text}")
                    if ocr_texts:
                        markdown_text = "\n\n".join(ocr_texts)
                except Exception as e:
                    print(f"[data_loader] PDF OCR Fallback failed: {e}")

    except Exception as e:
        raise ValueError(f"Lỗi khi parse file '{path_obj.name}': {e}")

    if not markdown_text.strip():
        raise ValueError(f"File sau khi parse rỗng (ngay cả khi đã thử OCR): '{path_obj.name}'.")

    # Bước 2: Dùng SentenceSplitter của LlamaIndex chia chunk dựa trên Markdown text
    splitter = SentenceSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    result = []
    
    doc = Document(
        id_=path_obj.name,
        text_resource=MediaResource(text=markdown_text),
        metadata={"source": path_obj.name, "page": 1, "type": file_type},
    )
    
    for node in splitter.get_nodes_from_documents([doc]):
        result.append({
            "text": node.get_content(),
            "page": 1,
            "type": file_type,
        })

    if not result:
        raise ValueError(f"Không tạo được chunk từ '{path_obj.name}'.")

    print(f"[data_loader] '{path_obj.name}' → {len(result)} chunks")
    return result


def embed_texts(texts: list[str], is_query: bool = False) -> list[list[float]]:
    """
    Embed list text. Ưu tiên dùng multilingual-e5-large (local).
    Nếu môi trường ML không sẵn sàng (lỗi import/compat) → fallback deterministic embedding
    để hệ thống vẫn chạy được các luồng tìm kiếm/caching cơ bản.

    Prefix:
    - "passage: " cho documents (khi index vào Qdrant)
    - "query: "   cho câu hỏi  (khi search)
    """
    import math, hashlib, random
    prefix = "query: " if is_query else "passage: "

    # Cho phép ép dùng dummy embed qua biến môi trường
    use_dummy = os.getenv("USE_DUMMY_EMBED", "0").strip() in ("1", "true", "True")

    global _model
    if not use_dummy and _model is None:
        try:
            # Import lazily to avoid expensive model load during module import time.
            from sentence_transformers import SentenceTransformer
            _model = SentenceTransformer(_MODEL_NAME)
        except Exception as e:
            print(f"[embed_texts] WARNING: sentence-transformers load failed: {e} → using dummy embeddings")
            _model = None
            use_dummy = True

    if not use_dummy and _model is not None:
        try:
            return _model.encode([prefix + t for t in texts], normalize_embeddings=True).tolist()
        except Exception as e:
            print(f"[embed_texts] WARNING: encode failed: {e} → falling back to dummy embeddings")
            use_dummy = True

    # Dummy embedding: deterministic hash-based 1024-dim unit vectors
    def _hash_vec(s: str, dim: int = 1024) -> list[float]:
        h = hashlib.md5(s.encode("utf-8")).digest()
        seed = int.from_bytes(h, "big")
        rnd = random.Random(seed)
        vec = [rnd.random() - 0.5 for _ in range(dim)]
        norm = math.sqrt(sum(x * x for x in vec)) or 1.0
        return [x / norm for x in vec]

    return [_hash_vec(prefix + t) for t in texts]
